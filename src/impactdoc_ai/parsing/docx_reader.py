"""DOCX dokümanlarını okuma işlemleri."""

from pathlib import Path

from docx import Document as DocxDocument

from impactdoc_ai.models.document import Document
from impactdoc_ai.utils.hashing import calculate_sha256


def read_docx(file_path: str | Path) -> Document:
    """
    Bir DOCX dosyasını okuyarak standart Document nesnesine dönüştürür.

    Args:
        file_path: Okunacak DOCX dosyasının yolu.

    Returns:
        Okunan dokümanı temsil eden Document nesnesi.

    Raises:
        FileNotFoundError: Dosya bulunamazsa.
        ValueError: Dosya uzantısı DOCX değilse veya metin çıkarılamazsa.
    """
    path = Path(file_path).resolve()

    if not path.exists():
        raise FileNotFoundError(f"Dosya bulunamadı: {path}")

    if not path.is_file():
        raise ValueError(f"Verilen yol bir dosya değil: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Desteklenmeyen dosya uzantısı: {path.suffix}. "
            "Bu okuyucu yalnızca DOCX dosyalarını destekler."
        )

    docx_document = DocxDocument(path)

    text_parts: list[str] = []

    for paragraph in docx_document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in docx_document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_values:
                text_parts.append(" | ".join(row_values))

    text = "\n".join(text_parts).strip()

    if not text:
        raise ValueError(
            "DOCX dosyasından okunabilir metin çıkarılamadı."
        )

    return Document(
        file_path=path,
        file_name=path.name,
        extension=path.suffix.lower(),
        text=text,
        page_count=1,
        character_count=len(text),
        word_count=len(text.split()),
        sha256=calculate_sha256(path),
    )